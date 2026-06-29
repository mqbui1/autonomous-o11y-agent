"""
Generate a DOCX assessment report for astronomy-shop-demo.
Run: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colors ────────────────────────────────────────────────────────────────────
SPLUNK_ORANGE   = RGBColor(0xF6, 0x6B, 0x00)
SPLUNK_DARK     = RGBColor(0x1A, 0x1A, 0x2E)
SPLUNK_GRAY     = RGBColor(0x60, 0x60, 0x60)
RED             = RGBColor(0xC0, 0x39, 0x2B)
ORANGE          = RGBColor(0xE6, 0x7E, 0x22)
YELLOW          = RGBColor(0xF3, 0x9C, 0x12)
GREEN           = RGBColor(0x27, 0xAE, 0x60)
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_RED_BG    = RGBColor(0xFF, 0xEB, 0xEB)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D0D0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_paragraph(doc, text="", bold=False, italic=False, size=11,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = SPLUNK_DARK
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    return p


def make_table(doc, headers, col_widths=None):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Inches(w)
    hrow = table.add_row()
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_cell_bg(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, "444444")
    return table


def add_table_row(table, cells, bold=False, bg=None, text_color=None, font_size=10):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        if text_color:
            run.font.color.rgb = text_color
        if bg:
            set_cell_bg(cell, bg)
        set_cell_borders(cell)
    return row


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "E0E0E0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

for txt, sz, color in [
    ("Autonomous Observability Agent", 28, SPLUNK_DARK),
    ("Assessment Report",              22, SPLUNK_ORANGE),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(sz)
    r.font.color.rgb = color

doc.add_paragraph()

for txt, sz in [
    ("Environment: astronomy-shop-demo", 14),
    ("Realm: us1  |  Mode: Dry Run  |  Platform: Splunk Observability Cloud", 12),
    (f"Generated: {datetime.datetime.utcnow().strftime('%B %d, %Y')}", 12),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.font.color.rgb = SPLUNK_GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Overall Environment Health Grade")
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = SPLUNK_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("D")
r.font.size = Pt(72); r.bold = True; r.font.color.rgb = RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Active PCI/PII data exposure + zero detector coverage + broken correlation links. "
    "Fixable to B+ within one sprint."
)
r.font.size = Pt(11); r.italic = True; r.font.color.rgb = SPLUNK_GRAY

doc.add_page_break()

# ── 1. AGENT OVERVIEW ─────────────────────────────────────────────────────────
add_heading(doc, "1. About the Autonomous O11y Agent", level=1, color=SPLUNK_DARK)
add_paragraph(doc,
    "The Autonomous Observability Agent is a multi-agent AI system built on AWS Bedrock "
    "(Claude Sonnet 4.6) that autonomously assesses, governs, and improves observability "
    "coverage for Splunk Observability Cloud environments. It replaces manual health checks "
    "with a continuous, data-driven loop that surfaces actionable, cross-domain findings.",
    size=11, space_after=6)

add_heading(doc, "Architecture", level=2, color=SPLUNK_DARK)
add_paragraph(doc, "Four specialist agents run in parallel, coordinated by a synthesis layer:", size=11, space_after=4)
for name, desc in [
    ("Health Agent",           "Audits detector quality, APM service coverage, OTel Collector health, and license utilization"),
    ("Instrumentation Agent",  "Scores APM span, infrastructure metric, and log attribute coverage (0-100); maps every gap to broken UI features"),
    ("Governance Agent",       "Scans metric cardinality, detects slow-burn MTS growth, snapshots trace volume, generates OTel Collector YAML fixes"),
    ("Detector Agent",         "Discovers services, learns behavioral baselines from live telemetry, provisions or retunes detectors"),
]:
    add_bullet(doc, f" - {desc}", bold_prefix=name)

add_paragraph(doc,
    "A Coordinator agent performs cross-domain analysis on all four specialist findings "
    "and synthesizes them into a unified prioritized report — surfacing compounding issues "
    "no single specialist could identify alone.",
    size=11, space_before=6, space_after=6)

add_heading(doc, "This Assessment Run", level=2, color=SPLUNK_DARK)
t = make_table(doc, ["Parameter", "Value"], col_widths=[2.2, 4.0])
for k, v in [
    ("Environment",    "astronomy-shop-demo"),
    ("Realm",          "us1"),
    ("Mode",           "Dry Run (no changes applied)"),
    ("Run Date",       datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ("Specialists",    "Health / Instrumentation / Governance / Detector"),
    ("Tools Executed", "13 tools across 4 agents + synthesis pass with all tools"),
    ("State File",     "~/.o11y-agent/astronomy-shop-demo.json"),
    ("New in v2",      "Structured output, cross-domain analysis, feedback loop, parallel subprocesses"),
]:
    add_table_row(t, [k, v])

doc.add_paragraph()
doc.add_page_break()

# ── 2. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
add_heading(doc, "2. Executive Summary", level=1, color=SPLUNK_DARK)
add_paragraph(doc,
    "The astronomy-shop-demo environment has clean telemetry governance (zero cardinality "
    "explosions, healthy license headroom, 0% APM error rate) sitting beneath a critically "
    "broken observability stack. Three compounding failures dominate: active PCI/PII data "
    "exposure, zero detector coverage on all 15 services, and systemic attribute gaps that "
    "break every correlation link.",
    size=11, space_after=8)

t = make_table(doc, ["Domain", "Status", "Key Finding"], col_widths=[1.8, 1.1, 3.9])
exec_rows = [
    ("Sensitive Data / PCI",    "CRITICAL", "5 spans across 4 services contain PII/card data — payment/charge = active PCI DSS scope"),
    ("Detector Coverage",       "CRITICAL", "0 of 142 recommended detectors deployed — all 15 services completely dark"),
    ("Instrumentation Quality", "CRITICAL", "Score 21/100; deployment.environment 0% on spans; all 3 correlation links broken"),
    ("Silent Services",         "HIGH",     "5 services silent: frontend, currency, shipping, quote, email"),
    ("APM -> IM Correlation",   "CRITICAL", "host.name 0% on spans, 81.5% missing on metrics; k8s.* 0% everywhere"),
    ("APM -> Logs Correlation", "CRITICAL", "Zero logs ingested; trace_id/span_id/deployment.environment absent"),
    ("Health-Check Pollution",  "MEDIUM",   "~35,110 liveness-probe spans/day on checkout + product-catalog"),
    ("OTel Collector",          "MEDIUM",   "Single collector v0.154.2 — within 30-90 day deprecation window; SPOF"),
    ("Telemetry Governance",    "HEALTHY",  "Zero cardinality explosions across 1,480 metrics; 0% APM error rate"),
    ("License Utilization",     "HEALTHY",  "APM 15.8%, IM 4%, Custom Metrics 19.7% — RUM and Synthetics unused"),
]
STATUS_COLORS = {
    "CRITICAL": RED, "HIGH": ORANGE, "MEDIUM": YELLOW,
    "HEALTHY": GREEN, "INFO": SPLUNK_GRAY,
}
for domain, status, finding in exec_rows:
    row = t.add_row()
    for i, (txt, col_idx) in enumerate([(domain, 0), (status, 1), (finding, 2)]):
        cell = row.cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(txt)
        run.font.size = Pt(10)
        if col_idx == 1:
            run.bold = True
            run.font.color.rgb = STATUS_COLORS.get(status, SPLUNK_GRAY)
        set_cell_borders(cell)

doc.add_paragraph()
add_paragraph(doc,
    "Bottom line: Three actions — (1) deploy redactionprocessor to stop active PCI/PII "
    "ingestion, (2) add OTEL_RESOURCE_ATTRIBUTES=deployment.environment=astronomy-shop-demo "
    "to every pod, (3) run provision_detectors(auto_deploy=True) — together address the "
    "majority of critical findings and can all be executed today.",
    bold=True, size=11, space_after=10)

doc.add_page_break()

# ── 3. CRITICAL: SENSITIVE DATA EXPOSURE ──────────────────────────────────────
add_heading(doc, "3. CRITICAL — Sensitive Data / PCI Exposure", level=1, color=RED)
add_paragraph(doc,
    "Five spans across four services contain PII or potential PCI DSS-scoped data. "
    "The payment/charge span is an active PCI DSS scoping event. Security team engagement "
    "is required before restoring shipping or quote instrumentation.",
    bold=True, size=11, space_after=8)

t = make_table(doc,
    ["Service", "Span", "Risk Category", "Action"],
    col_widths=[1.2, 2.2, 1.4, 2.0])
for svc, span, risk, action in [
    ("payment",  "charge",                                      "PCI DSS — card data",    "Engage security; assess breach notification"),
    ("checkout", "prepareOrderItemsAndShippingQuoteFromCart",   "PII — order/customer",   "Deploy redactionprocessor immediately"),
    ("checkout", "oteldemo.CheckoutService/PlaceOrder",         "PII — order/customer",   "Deploy redactionprocessor immediately"),
    ("shipping", "POST /get-quote",                             "GDPR/CCPA — address PII","Do NOT restore until redaction in place"),
    ("quote",    "calculate-quote",                             "Business-sensitive data", "Do NOT restore until redaction in place"),
]:
    add_table_row(t, [svc, span, risk, action])

doc.add_paragraph()
add_heading(doc, "Remediation — OTel Collector redactionprocessor", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""processors:
  redaction:
    allow_all_keys: true
    blocked_values:
      - "4[0-9]{12}(?:[0-9]{3})?"                              # Visa
      - "(?:5[1-5][0-9]{2}|222[1-9]|22[3-9][0-9]|2[3-6][0-9]{2}|27[01][0-9]|2720)[0-9]{12}"  # Mastercard
      - "[0-9]{3}-[0-9]{2}-[0-9]{4}"                          # SSN
      - "[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}"    # Email
    summary: debug

service:
  pipelines:
    traces:
      processors: [memory_limiter, redaction, resourcedetection, resource, k8sattributes, batch]""")

doc.add_page_break()

# ── 4. CROSS-DOMAIN ISSUES ────────────────────────────────────────────────────
add_heading(doc, "4. Cross-Domain Issues", level=1, color=SPLUNK_DARK)
add_paragraph(doc,
    "The following issues appear across two or more specialist domains and represent the "
    "highest-priority synthesis findings. They cannot be identified by any single specialist alone.",
    size=11, space_after=8)

cross_domain_issues = [
    ("checkout",
     "CRITICAL",
     "PII data in 2 spans + zero detectors + 17,560 health-check spans/day + "
     "deployment.environment absent. Revenue-critical service with active data leak, "
     "zero alerting, and no infrastructure correlation.",
     "Fix: redactionprocessor -> deploy 14 detectors -> add deployment.environment"),
    ("payment",
     "CRITICAL",
     "Active PCI DSS scoping event (charge span) + zero detectors for financial transactions "
     "(1.8 req/min) + no host.name or deployment.environment — no infrastructure or log "
     "correlation for PCI-scoped flows.",
     "Fix: Security team engagement -> breach assessment -> redactionprocessor -> detectors"),
    ("shipping + quote",
     "CRITICAL",
     "Both services generated sensitive-data findings then went completely silent. Most "
     "dangerous state: PII was ingested, the services are now dark, and there is no telemetry "
     "to determine if they are down or silently misbehaving.",
     "Fix: Do NOT restore instrumentation until redactionprocessor is deployed"),
    ("frontend",
     "HIGH",
     "Governance scan saw 295 spans/hour; health agent sees zero traces in 24h. "
     "This discrepancy indicates a recent instrumentation failure or deployment event. "
     "1,426 req/min baseline with no detectors = possible active incident with zero signal.",
     "Fix: Triage pod OTel SDK init logs -> check OTEL_EXPORTER_OTLP_ENDPOINT -> deploy 16 detectors"),
    ("currency + email",
     "HIGH",
     "Both silent for 24h with no detectors. currency runs on every product page load "
     "(130 req/min baseline). email silence means every post-order confirmation fails invisibly.",
     "Fix: Investigate OTel SDK liveness -> deploy 11/5 detectors"),
    ("deployment.environment absent",
     "CRITICAL",
     "0% on spans (0/163), 69% missing on metrics. Breaks: Service Centric View scoping, "
     "APM->Logs RC, APM->IM RC, all 142 detector environment filters, IM environment dashboards.",
     "Fix: One env var on every pod — OTEL_RESOURCE_ATTRIBUTES=deployment.environment=astronomy-shop-demo"),
    ("frontend-proxy",
     "HIGH",
     "Owns 75.2% of all trace volume (2,328,752 traces/day), p99=1,845ms unmonitored. "
     "An outage here silently takes down the entire application. Zero detectors.",
     "Fix: Deploy 12 detectors -> evaluate tail-based sampling strategy"),
]

for svc, sev, desc, fix in cross_domain_issues:
    add_heading(doc, svc, level=2,
                color=RED if sev == "CRITICAL" else ORANGE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"[{sev}]  ")
    r.bold = True
    r.font.color.rgb = STATUS_COLORS[sev]
    r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)
    add_paragraph(doc, fix, italic=True, size=10, color=SPLUNK_GRAY, space_after=8)
    add_divider(doc)

doc.add_page_break()

# ── 5. DETAILED FINDINGS PER DOMAIN ──────────────────────────────────────────
add_heading(doc, "5. Detailed Findings", level=1, color=SPLUNK_DARK)

# 5.1 Detector Coverage
add_heading(doc, "5.1  Detector Coverage — CRITICAL", level=2, color=RED)
add_paragraph(doc,
    "Zero detectors are deployed for any of the 15 astronomy-shop services. "
    "142 detectors are baselined and ready to deploy.",
    size=11, space_after=6)

t = make_table(doc,
    ["Service", "Stack", "Req/Min", "p99", "Detectors", "Notes"],
    col_widths=[1.3, 1.3, 0.7, 0.8, 0.85, 2.0])
for row in [
    ("frontend",        "Node.js + Next.js", "1,426",  "252.9ms",  "16", "Silent in 24h — triage first"),
    ("ad",              "JVM + gRPC",         "36.7",   "3.8ms",    "17", "JVM heap/GC/Tomcat coverage"),
    ("recommendation",  "Python + gRPC",      "108",    "4.4ms",    "14", "GC, thread, RSS"),
    ("product-catalog", "Go + gRPC",          "642",    "4.6ms",    "14", "Health-probe noise; goroutine leak"),
    ("cart",            ".NET + Redis",        "—",      "—",        "20", "Highest detector count"),
    ("product-reviews", "Python + PostgreSQL", "~37",    "11.3s LLM","17", "ONLY GenAI service — LLM p99=11.3s"),
    ("payment",         "Node.js + gRPC",      "1.8",    "—",        "7",  "PCI scope — deploy ASAP"),
    ("checkout",        "Go + gRPC",           "13.2",   "—",        "14", "PII leak; health-probe noise"),
    ("currency",        "gRPC",                "130.3",  "0.3ms",    "11", "Silent in 24h; every page load"),
    ("frontend-proxy",  "Istio",               "1,479",  "1,845ms",  "12", "75.2% of trace volume; SPOF risk"),
    ("shipping",        "Rust",                "20",     "3.6ms",    "6",  "Silent + PII — do not restore yet"),
    ("email",           "APM",                 "1.8",    "—",        "5",  "Silent in 24h"),
    ("quote",           "APM",                 "—",      "—",        "5",  "Silent + PII — do not restore yet"),
    ("image-provider",  "APM",                 "—",      "—",        "5",  ""),
    ("telemetry-docs",  "APM",                 "—",      "—",        "5",  "Orphan — no env tag"),
    ("TOTAL",           "",                    "",       "",         "142",""),
]:
    is_total = row[0] == "TOTAL"
    add_table_row(t, row, bold=is_total,
                  bg=TABLE_HEADER_BG if is_total else None,
                  text_color=WHITE if is_total else None)

add_paragraph(doc,
    "product-reviews is the only GenAI/LLM service in the fleet — LLM operation p99=11.3s. "
    "Recommended thresholds: warn >15s, critical >30s (1.5x/2.5x observed p99). "
    "frontend-proxy p99=1,845ms is already elevated — investigate before deployment "
    "to avoid alert storms on a pre-existing anomaly.",
    italic=True, size=10, color=ORANGE, space_before=6, space_after=10)

# 5.2 Instrumentation Quality
add_heading(doc, "5.2  Instrumentation Quality — CRITICAL (Score: 21/100)", level=2, color=RED)

t = make_table(doc, ["Signal", "Score", "Status"], col_widths=[2.5, 1.0, 3.2])
for sig, score, status in [
    ("APM Spans",              "21/100", "Critical — all correlation attrs missing"),
    ("Infrastructure Metrics", "22/100", "Critical — environment/host/k8s gaps"),
    ("Logs",                   "0/100",  "No data — pipeline not configured"),
    ("Combined",               "21/100", "All 3 Related Content links broken"),
]:
    add_table_row(t, [sig, score, status], bold=(sig == "Combined"))

doc.add_paragraph()
t2 = make_table(doc,
    ["Missing Attribute", "Coverage Gap", "UX Impact"],
    col_widths=[1.8, 1.6, 3.0])
for attr, gap, impact in [
    ("deployment.environment", "0% spans / 69% metrics", "SCV blank; APM->IM and APM->Logs broken; detector scoping unreliable"),
    ("host.name",              "0% spans / 81.5% metrics","APM->IM join broken; Host Navigator empty"),
    ("k8s.pod.name",           "0% everywhere",           "K8s Navigator completely dark"),
    ("k8s.node.name",          "0% everywhere",           "Node rollup broken"),
    ("k8s.namespace.name",     "0% everywhere",           "Namespace drill-down broken"),
    ("container.id",           "0% spans",                "Container-level IM correlation broken"),
    ("telemetry.sdk.*",        "0% spans",                "SDK version alerts and language dashboards non-functional"),
    ("Logs (all)",             "0 records ingested",      "Log Observer non-functional; APM->Logs impossible"),
    ("Runtime metrics",        "0% all runtimes",         "Runtime tab empty for JVM/.NET/Node.js/Python services"),
    ("http.status_code",       "42.3% coverage",          "57.7% of HTTP errors not classified by status code"),
]:
    add_table_row(t2, [attr, gap, impact])

# 5.3 APM Health
doc.add_paragraph()
add_heading(doc, "5.3  APM Service Health", level=2, color=ORANGE)

t3 = make_table(doc,
    ["Service", "Traces (24h)", "% of Total", "Note"],
    col_widths=[1.5, 1.3, 1.0, 3.0])
for svc, traces, pct, note in [
    ("frontend-proxy",  "2,328,752", "75.2%",  "Ingress; p99=1,845ms — evaluate tail sampling"),
    ("frontend-web",    "611,152",   "19.7%",  "Active"),
    ("load-generator",  "~3,995",    "~0.5%",  "Synthetic load"),
    ("product-catalog", "~17,550",   "<0.1%",  "Health-probe noise dominates"),
    ("checkout",        "~17,560",   "<0.1%",  "Health-probe noise dominates + PII spans"),
    ("image-provider",  "~854",      "<0.1%",  "Active"),
    ("telemetry-docs",  "~770",      "<0.1%",  "Orphan — no deployment.environment"),
    ("cart",            "~167",      "<0.1%",  "Active"),
    ("ad",              "~18",       "<0.1%",  "Active"),
    ("product-reviews", "~16",       "<0.1%",  "Active — LLM p99=11.3s"),
    ("payment",         "~6",        "<0.1%",  "Active — PCI scope"),
    ("frontend",        "0",         "—",      "SILENT — 295 spans/hr seen in governance scan (recent failure)"),
    ("currency",        "0",         "—",      "SILENT — had 130 req/min baseline"),
    ("shipping",        "0",         "—",      "SILENT + PII spans exist — do not restore yet"),
    ("quote",           "0",         "—",      "SILENT + PII spans exist — do not restore yet"),
    ("email",           "0",         "—",      "SILENT — had 1.8 req/min baseline"),
]:
    is_silent = "SILENT" in note
    add_table_row(t3, [svc, traces, pct, note],
                  bg=LIGHT_RED_BG if is_silent else None)

doc.add_paragraph()
for title, detail in [
    ("Health-check span pollution",
     "checkout + product-catalog each emit ~17,550 grpc.health.v1.Health/Check spans/day "
     "(~35,110 total). 100% synthetic noise. Drop at Collector or migrate to Synthetics."),
    ("frontend vs governance discrepancy",
     "Governance scan (1h window) saw frontend at 295 spans/hr. Health scan (24h window) "
     "sees zero. This temporal gap strongly suggests a recent deployment event or "
     "instrumentation failure — treat as active incident until confirmed otherwise."),
    ("db.query.text indexing",
     "36,388 chars indexed. Full SQL with potential PII in WHERE clauses. "
     "High cardinality risk. Normalize or remove from indexed scope."),
    ("OTel Collector",
     "Single instance bd884ccbf517 on v0.154.2 (Yellow status). "
     "30-90 day deprecation window. Single point of failure for all pipelines."),
    ("Orphan: telemetry-docs",
     "No deployment.environment, no upstream/downstream dependencies. Set tag or decommission."),
]:
    add_bullet(doc, f": {detail}", bold_prefix=title)

# 5.4 Governance
doc.add_paragraph()
add_heading(doc, "5.4  Telemetry Governance — HEALTHY", level=2, color=GREEN)
t4 = make_table(doc, ["Metric", "Value"], col_widths=[3.0, 3.8])
for k, v in [
    ("Total metrics scanned",         "1,480"),
    ("Cardinality explosions",        "0"),
    ("Anomaly flags",                 "0 (baseline not yet established — requires 3+ daily scans)"),
    ("Active APM services (1h scan)", "19"),
    ("Global APM error rate",         "0.0%"),
    ("Total spans (1h)",              "1,103"),
]:
    add_table_row(t4, [k, v])

# 5.5 License
doc.add_paragraph()
add_heading(doc, "5.5  License Utilization — HEALTHY", level=2, color=GREEN)
t5 = make_table(doc, ["Dimension", "Utilization", "Notes"], col_widths=[2.2, 1.3, 3.3])
for dim, util, note in [
    ("APM Hosts",        "15.84%",  "Healthy"),
    ("APM Containers",   "0.16%",   "Significant headroom"),
    ("IM Hosts",         "4.0%",    "Healthy"),
    ("Custom Metrics",   "19.68%",  "Monitor after k8s attr addition"),
    ("RUM Sessions",     "0%",      "Entitlement active — no RUM configured"),
    ("Synthetics",       "0%",      "Entitlement active — ideal for replacing health-probe spans"),
]:
    add_table_row(t5, [dim, util, note])

doc.add_page_break()

# ── 6. PRIORITIZED ACTION PLAN ────────────────────────────────────────────────
add_heading(doc, "6. Prioritized Action Plan", level=1, color=SPLUNK_DARK)

add_heading(doc, "Immediate — Do Today", level=2, color=RED)
t6 = make_table(doc, ["#", "Action", "Impact", "Effort"], col_widths=[0.35, 3.1, 2.2, 0.9])
for num, action, impact, effort in [
    ("I-1", "Deploy redactionprocessor — stop PCI/PII ingestion on payment, checkout, shipping, quote",
             "Stops active data breach; required before restoring silent services", "Low"),
    ("I-2", "Set deployment.environment=astronomy-shop-demo on every pod (one env var)",
             "Unblocks all 3 RC links; fixes SCV; enables 142 detector scoping", "Low"),
    ("I-3", "Run provision_detectors(auto_deploy=True) — deploy all 142 detectors",
             "Eliminates complete alerting blindness across all 15 services", "Low"),
    ("I-4", "Triage 5 silent services (frontend priority — recent failure pattern)",
             "frontend was active 1h before health scan — likely active incident", "Medium"),
]:
    row = t6.add_row()
    for i, txt in enumerate([num, action, impact, effort]):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(txt)
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True; r.font.color.rgb = RED
        set_cell_borders(cell)

doc.add_paragraph()
add_heading(doc, "Short-Term — 1-2 Sprints", level=2, color=ORANGE)
t7 = make_table(doc, ["#", "Action", "Impact", "Effort"], col_widths=[0.35, 3.1, 2.2, 0.9])
for num, action, impact, effort in [
    ("S-1", "Add k8sattributes processor + resourcedetection to all pipelines",
             "Restores APM->IM; populates K8s Navigator; adds host.name", "Medium"),
    ("S-2", "Configure log ingestion + inject trace_id/span_id/deployment.environment",
             "Enables APM->Logs RC; log score 0->functional", "Medium"),
    ("S-3", "Drop grpc.health.v1.Health/Check spans via filter processor",
             "Eliminates ~35,110 synthetic spans/day; reduces APM quota waste", "Low"),
    ("S-4", "Enable runtime metrics for JVM/(.NET)/Node.js/Python/Go services",
             "Populates Runtime tab in SCV for all 6 active language runtimes", "Medium"),
    ("S-5", "Upgrade OTel Collector from v0.154.2 to latest stable",
             "Exits deprecation; unlocks latest processors; resolves SPOF concern", "Low"),
    ("S-6", "Schedule daily full_cardinality_scan cron job (06:00 UTC)",
             "Builds anomaly baseline; active slow-burn detection in 7 days", "Low"),
    ("S-7", "Apply tail-based sampling on frontend-proxy (75.2% of trace volume)",
             "Reduces APM quota while preserving error/slow traces", "Medium"),
    ("S-8", "Standardize service names (postgresql, flagd variants)",
             "Fixes service map fragmentation; improves dashboard grouping", "Low"),
]:
    row = t7.add_row()
    for i, txt in enumerate([num, action, impact, effort]):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(txt)
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True; r.font.color.rgb = ORANGE
        set_cell_borders(cell)

doc.add_paragraph()
add_heading(doc, "Ongoing", level=2, color=SPLUNK_GRAY)
for num, action in [
    ("O-1",  "Run retune_detectors after 7 days production traffic for 8 thin-baseline services"),
    ("O-2",  "Re-run provision_detectors(baseline_window_hours=336) after 14 days of history"),
    ("O-3",  "Configure Synthetics uptime checks to replace gRPC health probes (activates unused entitlement)"),
    ("O-4",  "Evaluate RUM activation for frontend-web (0% utilization with active entitlement)"),
    ("O-5",  "Audit db.query.text indexed tag — truncate or remove from indexing scope"),
    ("O-6",  "Add cloud resource detectors (cloud.provider, cloud.region, cloud.account.id)"),
    ("O-7",  "Investigate frontend-proxy p99=1,845ms baseline anomaly before first detector retune"),
    ("O-8",  "Resolve telemetry-docs orphan — set deployment.environment or decommission"),
    ("O-9",  "Review product-reviews LLM thresholds after 7-day production baseline"),
    ("O-10", "Add LLM-specific detectors for token usage and model latency on product-reviews"),
]:
    add_bullet(doc, f" {action}", bold_prefix=num)

doc.add_page_break()

# ── 7. AGENT CAPABILITIES & GAPS ─────────────────────────────────────────────
add_heading(doc, "7. Agent Capabilities and Known Gaps", level=1, color=SPLUNK_DARK)

t8 = make_table(doc, ["Capability", "Status", "Notes"], col_widths=[2.4, 1.0, 3.4])
for cap, status, notes in [
    ("Detector health audit",              "Live", "Ghost/noisy/muted/never-fired checks; env-scoped filtering"),
    ("APM service health",                 "Live", "Silent services, health-probe noise, sensitive data, orphans"),
    ("Instrumentation quality scoring",    "Live", "0-100 per signal; maps gaps to specific RC link failures"),
    ("Metric cardinality governance",      "Live", "MTS explosion detection + YAML fix generation"),
    ("Parallel cardinality+anomaly scan",  "Live", "batch_run() runs both subprocesses concurrently (Gap 8)"),
    ("Detector provisioning",              "Live", "30+ stacks; dynamic baseline learning from live telemetry"),
    ("GenAI/LLM service detection",        "Live", "product-reviews auto-detected; LLM-specific detector templates"),
    ("Parallel multi-agent execution",     "Live", "4 specialists concurrent; ~5 min vs ~17 min sequential"),
    ("Structured output (Gap 6)",          "Live", "submit_findings schema; SpecialistFindings dataclass"),
    ("Cross-domain analysis (Gap 4)",      "Live", "Post-parallel service/issue correlation across all 4 domains"),
    ("Synthesis with tools (Gap 5)",       "Live", "Synthesis LLM has all 13 tools for targeted follow-up"),
    ("Rich persistent state (Gap 3)",      "Live", "Typed RunRecord: silent names, deployed IDs, critical issues"),
    ("Feedback loop (Gap 7)",              "Live", "Prior silent services + detector IDs injected on next run"),
    ("Human-in-the-loop approval",         "Planned", "Auto-apply is binary; no tiered approval or Slack notification"),
    ("Alerting / notification output",     "Planned", "Findings go to stdout only; no Slack/PD/email routing"),
    ("Agent self-observability",           "Planned", "No OTel instrumentation on the agent itself"),
    ("RUM analysis",                       "Planned", "No RUM assessment tool"),
    ("SLO management",                     "Planned", "Cannot create/read/modify SLOs"),
    ("Multi-environment fan-out",          "Planned", "One environment per run; no cross-env comparison"),
]:
    is_live = status == "Live"
    add_table_row(t8, [cap, status, notes],
                  text_color=GREEN if is_live else SPLUNK_GRAY)

doc.add_page_break()

# ── APPENDIX A — OTel Config Recipes ─────────────────────────────────────────
add_heading(doc, "Appendix A — OTel Collector Configuration Recipes", level=1, color=SPLUNK_DARK)

add_heading(doc, "A.1  Add deployment.environment + resourcedetection (I-2 + S-1)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""processors:
  resource:
    attributes:
      - key: deployment.environment
        value: "astronomy-shop-demo"
        action: upsert
      - key: sf_environment
        value: "astronomy-shop-demo"
        action: upsert
  resourcedetection:
    detectors: [system, k8snode, docker, env]
    system:
      hostname_sources: ["os"]
    override: false
service:
  pipelines:
    traces:
      processors: [memory_limiter, redaction, resourcedetection, resource, k8sattributes, batch]
    metrics:
      processors: [memory_limiter, resourcedetection, resource, k8sattributes, batch]
    logs:
      processors: [memory_limiter, resourcedetection, resource, k8sattributes, batch]""")

add_heading(doc, "A.2  redactionprocessor — stop PCI/PII ingestion (I-1)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""processors:
  redaction:
    allow_all_keys: true
    blocked_values:
      - "4[0-9]{12}(?:[0-9]{3})?"
      - "(?:5[1-5][0-9]{2}|222[1-9]|22[3-9][0-9]|2[3-6][0-9]{2}|27[01][0-9]|2720)[0-9]{12}"
      - "[0-9]{3}-[0-9]{2}-[0-9]{4}"
      - "[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}"
    summary: debug""")

add_heading(doc, "A.3  k8sattributes processor (S-1)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""processors:
  k8sattributes:
    auth_type: "serviceAccount"
    extract:
      metadata:
        - k8s.pod.name
        - k8s.pod.uid
        - k8s.node.name
        - k8s.namespace.name
        - k8s.deployment.name
    pod_association:
      - sources:
          - from: resource_attribute
            name: k8s.pod.ip
      - sources:
          - from: connection""")

add_heading(doc, "A.4  Drop health-check spans (S-3)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""processors:
  filter/drop_health_checks:
    traces:
      span:
        - >
          attributes["rpc.method"] == "Check" and
          attributes["rpc.service"] == "grpc.health.v1.Health"
service:
  pipelines:
    traces:
      processors: [memory_limiter, redaction, filter/drop_health_checks,
                   resourcedetection, resource, k8sattributes, batch]""")

add_heading(doc, "A.5  Runtime metrics — JVM (S-4)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""# Environment variables for ad service
OTEL_INSTRUMENTATION_RUNTIME_METRICS_ENABLED=true
SPLUNK_METRICS_ENABLED=true
OTEL_EXPORTER_OTLP_ENDPOINT=http://<collector>:4317""")

add_heading(doc, "A.6  Daily cardinality scan cron (S-6)", level=2, color=SPLUNK_DARK)
add_code_block(doc,
"""# crontab -e
# Daily at 06:00 UTC — builds anomaly baseline; active in 7 days
0 6 * * * cd /path/to/o11y-usage-governance && \\
  python3 cardinality_governance.py scan >> /var/log/cardinality.log 2>&1""")

# ── APPENDIX B — Re-running ───────────────────────────────────────────────────
add_heading(doc, "Appendix B — Re-running the Agent", level=1, color=SPLUNK_DARK)
add_code_block(doc,
"""# Dry run
python3 main.py --realm us1 --token $TOKEN --environment astronomy-shop-demo

# Deploy all detectors
python3 main.py --realm us1 --token $TOKEN --environment astronomy-shop-demo --auto-apply

# Scope to one service
python3 main.py --realm us1 --token $TOKEN --environment astronomy-shop-demo \\
  --service currency

# Ask a targeted question
python3 main.py --realm us1 --token $TOKEN --environment astronomy-shop-demo \\
  --prompt "Why is currency silent and what do I need to fix?"

# Watch mode — every 30 minutes
python3 main.py --realm us1 --token $TOKEN --environment astronomy-shop-demo \\
  --watch --interval 30""")

add_paragraph(doc,
    "State is persisted at ~/.o11y-agent/astronomy-shop-demo.json. "
    "Each subsequent run receives trend context including prior silent services and deployed "
    "detector IDs, enabling explicit regression detection and feedback-loop verification.",
    size=10, italic=True, color=SPLUNK_GRAY)

# ── SAVE ──────────────────────────────────────────────────────────────────────
output_path = "/Users/mbui/Documents/autonomous-o11y-agent/AstronomyShop_O11y_Assessment.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
