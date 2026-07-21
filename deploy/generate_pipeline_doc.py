#!/usr/bin/env python3
"""Generate the pipeline overview Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xCC)
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xCC)
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.size = Pt(12)
    return p

def para(text="", bold=False, italic=False, size=10.5, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F4FF")
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A56CC")
        tc_pr.append(shd)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for row_idx, row in enumerate(rows):
        r = t.add_row()
        fill = "F5F8FF" if row_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Autonomous O11y Agent")
run.font.size  = Pt(26)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xCC)

p2 = doc.add_paragraph()
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Assessment → Labeling → Fine-Tuning Pipeline")
run2.font.size  = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p3 = doc.add_paragraph()
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Technical Overview for Stakeholders")
run3.font.size  = Pt(11)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# Status box
status_p = doc.add_paragraph()
status_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = status_p.add_run(
    "Current Status  |  Phase 1 of 5  |  370 labeled examples  |  86% approval rate"
)
sr.font.size  = Pt(10)
sr.font.bold  = True
sr.font.color.rgb = RGBColor(0x0A, 0x7A, 0x3A)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

heading("1. What This Is", level=1)
para(
    "The Autonomous O11y Agent is a self-contained training data flywheel. It continuously "
    "runs AI-powered observability assessments against a live Splunk Observability Cloud "
    "environment, automatically labels the results using deterministic rules, and builds "
    "a structured JSONL dataset across multiple environment states. At the end of the "
    "pipeline, the dataset is used to fine-tune a local LLM — teaching it what good "
    "and bad observability analysis looks like in practice."
)
para(
    "No human labeling is required. The entire pipeline — assessment, labeling, phase "
    "transitions, environment fixes, export, and fine-tuning — runs unattended."
)

doc.add_paragraph()
heading("Key Design Principles", level=2)
bullet("Rule-based labeling over LLM-as-judge — deterministic, instant, zero extra API cost")
bullet("Phase curriculum — broken → fixed → high-load → RUM-active states create richer training contrast than repeated identical runs")
bullet("Reject filtering — tooling failures and credential errors produce < 150-char outputs; these are auto-rejected and never included in training data")
bullet("State persistence — pipeline_state.json tracks current phase so restarts resume exactly where they left off")
bullet("Fully unattended — no human intervention needed once started")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

heading("2. System Architecture", level=1)

code_block(
    "OTel Demo (Astroshop, 20+ services)\n"
    "    │  traces + metrics\n"
    "    ▼\n"
    "Splunk OTel Collector  ──►  Splunk Observability Cloud (us1)\n"
    "                                       │\n"
    "                                       │  Splunk APIs\n"
    "                                       ▼\n"
    "                            o11y-agent  (10 AI specialists)\n"
    "                                       │  assessment results\n"
    "                                       ▼\n"
    "                            Supervisor  (FastAPI + Web UI)\n"
    "                             ▲         │  train.jsonl\n"
    "                             │         ▼\n"
    "                       auto_labeler  Fine-tuning (MLX)\n"
    "                       (approve/reject rules)\n"
    "\n"
    "Orchestration: training_pipeline.py\n"
    "  — monitors labeled counts, advances phases, applies fixes"
)

heading("Components", level=2)

add_table(
    ["Component", "Technology", "Role"],
    [
        ["OTel Demo (Astroshop)", "OpenTelemetry demo, 20+ microservices", "Generates live traces, metrics, spans to Splunk"],
        ["Splunk OTel Collector", "OpenTelemetry Collector with Splunk exporters", "Receives telemetry, applies processors (filter, redaction), forwards to Splunk"],
        ["o11y-agent", "Python, Claude Sonnet 4.6 via AWS Bedrock", "Runs 10 parallel specialist agents per assessment cycle"],
        ["Supervisor", "FastAPI, Python", "Stores assessment results, labels, exposes REST API and Web UI"],
        ["auto_labeler.py", "Python, rule-based", "Polls for new runs, applies approve/reject labels per domain"],
        ["training_pipeline.py", "Python, orchestrator", "Monitors phase progress, applies environment fixes, triggers exports"],
        ["browser-sim", "Playwright (headless Chromium)", "Simulates real browser sessions for Splunk RUM data generation"],
        ["finetune.py", "MLX, Apple Silicon MPS", "Runs fine-tuning on local model using exported train.jsonl"],
    ],
    col_widths=[1.8, 2.0, 3.2]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. STEP 1 — ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Step 1 — Assessment (o11y-agent)", level=1)
para(
    "Every 10 minutes the o11y-agent dispatches 10 parallel AI specialist agents, each "
    "querying a different Splunk API surface. Each specialist uses Claude Sonnet 4.6 "
    "(via AWS Bedrock) with tool-use to call Splunk APIs, process results, and produce "
    "structured JSON output. The full assessment is stored in the supervisor under a "
    "unique run_id."
)

heading("The 10 Specialists", level=2)
add_table(
    ["Specialist", "What It Analyzes", "Key Signals"],
    [
        ["health", "Service error rates, latency, top issues", "Error %, P99 latency, service rankings"],
        ["instrumentation", "OTel attribute completeness", "service.name coverage, missing attributes, K8s vs Docker env"],
        ["governance", "Cardinality, PII in spans, metric volume", "High-cardinality dimensions, private IPs in attributes"],
        ["detector", "Alert detector coverage and correctness", "Missing detectors, cross-environment contamination"],
        ["rca", "Root cause analysis for top errors", "Trace-level error chains, downstream impact"],
        ["logs", "Log Observer availability and quality", "HTTP 404/403 status, log volume, error patterns"],
        ["rum", "Real User Monitoring sessions", "Session count, Core Web Vitals (LCP/FID/CLS), JS errors"],
        ["synthetics", "Synthetic test coverage", "HTTP 403 no-entitlement, test gap analysis"],
        ["db", "Database span attributes", "db.* attribute presence, query visibility, flagd error rate"],
        ["performance", "Profiling coverage", "CPU/memory hotspots, profiling gaps"],
    ],
    col_widths=[1.4, 2.2, 3.4]
)

heading("Assessment Output", level=2)
para("Each specialist produces a structured result stored per run_id:")
code_block(
    '{\n'
    '  "run_id": "run_65c4d4802b",\n'
    '  "instrumentation_score": 28,\n'
    '  "specialists": {\n'
    '    "governance": {\n'
    '      "raw_text": "Governance analysis found 3 PII issues...",\n'
    '      "issues": [\n'
    '        {"description": "Private IPv4 address in upstream_address",\n'
    '         "severity": "high",\n'
    '         "recommendation": "Enable PII redaction processor"}\n'
    '      ]\n'
    '    },\n'
    '    ... (9 more specialists)\n'
    '  }\n'
    '}'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. STEP 2 — AUTO-LABELING
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Step 2 — Auto-Labeling (auto_labeler.py)", level=1)
para(
    "The auto_labeler polls the supervisor every 60 seconds for unlabeled runs and applies "
    "deterministic approve/reject rules per domain. Rules encode known facts about the "
    "target environment — facts that a human reviewer would need significant context to apply "
    "consistently. By encoding them once, labeling is instant, free, and reproducible."
)

heading("Known Environment Facts (Astroshop Local)", level=2)
bullet("Runtime: Docker Compose (NOT Kubernetes)")
bullet("Log Observer: HTTP 404 on all log APIs — not licensed")
bullet("Synthetics: HTTP 403 — no Synthetics entitlement on this account")
bullet("RUM: No browser sessions Phase 0–2; Playwright browser-sim active from Phase 3")
bullet("Cross-environment risk: Splunk org contains detectors from other customers/envs (Nuveen, SOR-Data, etc.)")

heading("Labeling Rules by Domain", level=2)
add_table(
    ["Domain", "Approve if", "Reject if"],
    [
        ["instrumentation",
         "Findings consistent with observable environment state",
         "Claims service.name missing from 100% of spans while 5+ named services are visible (internal contradiction)\nOR 4+ K8s-specific recommendations for a Docker Compose environment"],
        ["detector",
         "Analysis correctly scoped to this environment",
         "Cross-environment contamination — detectors from other orgs appear in output (Nuveen, SOR-Data, petclinic)"],
        ["governance",
         "Accurately reports PII in-flight detections and/or tooling errors",
         "Returns zero findings (tooling completely offline)"],
        ["logs",
         "Correctly identifies Log Observer as unavailable (HTTP 404)",
         "—"],
        ["synthetics",
         "Correctly identifies Synthetics as unavailable (HTTP 403)",
         "—"],
        ["rum",
         "Zero sessions (Phase 0–2) or active session/vitals data (Phase 3+)",
         "—"],
        ["health / rca / db / performance",
         "Always approved — reliable streaming pipeline data",
         "—"],
        ["Any domain",
         "—",
         "Raw output < 150 characters (tooling failure, AWS credential expiry, or timeout)"],
    ],
    col_widths=[1.5, 2.8, 2.7]
)

heading("Label Output Format", level=2)
para("Labels are appended to tuning_decisions.jsonl in the supervisor's data volume:")
code_block(
    '{"run_id": "run_65c4d4802b", "domain": "governance", "decision": "approve", "timestamp": "2026-07-18T05:40:31Z"}\n'
    '{"run_id": "run_65c4d4802b", "domain": "instrumentation", "decision": "reject", "timestamp": "2026-07-18T05:40:31Z"}\n'
    '...'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. STEP 3 — TRAINING DATA EXPORT
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Step 3 — Training Data Export", level=1)
para(
    "The supervisor's /api/training/export endpoint merges assessment results with labels "
    "and streams a train.jsonl file. Only approved specialist outputs are included — "
    "rejected outputs are filtered out entirely so the model never trains on garbage data."
)
heading("JSONL Record Structure", level=2)
code_block(
    '{\n'
    '  "run_id": "run_65c4d4802b",\n'
    '  "domain": "governance",\n'
    '  "decision": "approve",\n'
    '  "specialist_output": {\n'
    '    "raw_text": "Governance analysis found 3 PII issues in span attributes...",\n'
    '    "issues": [...],\n'
    '    "recommendations": [...]\n'
    '  },\n'
    '  "environment_context": {\n'
    '    "phase": 1,\n'
    '    "instrumentation_score": 28,\n'
    '    "timestamp": "2026-07-18T05:40:29Z"\n'
    '  }\n'
    '}'
)
para(
    "Snapshots are exported at the end of Phase 1 and Phase 4. The Phase 4 export "
    "contains the complete multi-phase dataset used for fine-tuning."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. STEP 4 — TRAINING PIPELINE ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════

heading("6. Step 4 — Training Pipeline Orchestrator (training_pipeline.py)", level=1)
para(
    "The orchestrator runs alongside the auto_labeler and polls the supervisor every 60 "
    "seconds. When a phase's labeled run target is reached, it applies the configured "
    "fixes for the next phase, optionally exports a JSONL snapshot, and advances. "
    "State is persisted in pipeline_state.json so the process can be safely restarted."
)

heading("Why Phases? — Curriculum Learning", level=2)
para(
    "A model trained only on a fixed, healthy environment would learn ideal behavior "
    "but not how to recognize or handle broken states. The phase curriculum deliberately "
    "captures multiple environment conditions so the fine-tuned model understands both "
    "what good analysis looks like AND what broken analysis looks like — and why."
)
para(
    "The before/after contrast between Phase 0 (broken) and Phase 1 (fixed) on the same "
    "environment is particularly high-value: the model sees the same specialist running "
    "on the same infrastructure in two different states, with labeled examples showing "
    "what correct analysis looks like in each."
)

heading("The 5 Phases", level=2)
add_table(
    ["Phase", "Name", "Runs", "Environment State", "Training Signal"],
    [
        ["0", "baseline", "20",
         "Broken: governance SyntaxError, flagd trace noise (99.75% of volume), cross-env detector contamination",
         "What broken looks like — low scores (~16–28), tooling failures, contradictory findings"],
        ["1", "fix_governance_and_trace_noise", "20",
         "Fixed: governance works, flagd noise eliminated, PII redacted from spans",
         "Score jump to ~50–70. Correct analysis after fixes — governance, instrumentation, detector specialists recover"],
        ["2", "increase_traffic_load", "20",
         "High load: 15 concurrent users (up from 5)",
         "Richer error patterns, latency degradation visible, service rankings change under stress"],
        ["3", "enable_rum_sessions", "20",
         "RUM active: Playwright browser-sim generates real browser sessions via rum-proxy",
         "RUM specialist transitions from 'zero sessions' to Core Web Vitals, JS error analysis"],
        ["4", "final_export_and_finetune", "15",
         "All environment states captured",
         "Final dataset export + MLX fine-tune trigger"],
    ],
    col_widths=[0.5, 2.0, 0.6, 2.2, 2.2]
)

heading("Fixes Applied at Phase Boundaries", level=2)

para("Phase 0 → 1  (applied 09:50 UTC, July 18 — COMPLETED)", bold=True)
bullet("Patched cardinality_governance.py f-string SyntaxError (Python 3.11 incompatibility in agent container)")
bullet("Added filter/drop_flagd_stream processor to OTel Collector — drops /flagd.evaluation.v2.Service/EventStream spans")
bullet("Added redaction/pii_ipv4 processor to OTel Collector — redacts private IPv4s from upstream_address, peer.address, net.sock.peer.addr")
bullet("Rebuilt o11y-agent Docker container to pick up governance code fix")
bullet("Restarted otel-collector to load new processors")

para()
para("Phase 1 → 2", bold=True)
bullet("Sets LOCUST_USERS=15 in .env (up from 5)")
bullet("Restarts load-generator container")

para()
para("Phase 2 → 3", bold=True)
bullet("Builds browser-sim Docker image (Playwright/Chromium, ~3 min first build)")
bullet("Starts browser-sim container via docker compose --profile rum up -d")
bullet("Simulator visits Astroshop UI via rum-proxy, triggering Splunk RUM JS SDK")
bullet("Real browser sessions appear in Splunk Observability Cloud within ~2 minutes")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. STEP 5 — FINE-TUNING
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Step 5 — Fine-Tuning (finetune.py)", level=1)
para(
    "Triggered automatically when Phase 4 completes. Uses MLX fine-tuning on Apple Silicon "
    "(M-series MPS GPU). Input is the complete train.jsonl exported at Phase 4 completion."
)
bullet("Framework: MLX (Apple's ML framework for M-series chips)")
bullet("Hardware: Apple Silicon MPS — no cloud GPU required")
bullet("Input: train.jsonl (~75 runs × ~8 approved specialists avg = ~600 training examples)")
bullet("Output: fine-tuned model adapter weights saved locally")


# ══════════════════════════════════════════════════════════════════════════════
# 8. CURRENT STATUS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("8. Current Pipeline Status", level=1)

add_table(
    ["Metric", "Value"],
    [
        ["Current phase", "Phase 1 — fix_governance_and_trace_noise"],
        ["Phase progress", "17 / 20 labeled runs (3 to go)"],
        ["Total labeled examples", "370"],
        ["Approved", "318 (86%)"],
        ["Rejected", "52 (14%)"],
        ["Phase 0 approval rate", "75%  (broken environment — expected)"],
        ["Phase 1 approval rate", "86%  (fixes working — improvement confirmed)"],
        ["Phase 0 → 1 fixes applied", "09:50 UTC, July 18 — all 5 fixes successful"],
        ["Next milestone", "Phase 1 complete → Phase 2 fixes (~15:50 UTC)"],
        ["Estimated completion", "Phase 4 complete ~01:00 UTC, July 19 (~6:00 PM PDT)"],
    ],
    col_widths=[2.8, 4.2]
)

heading("Remaining Schedule", level=2)
add_table(
    ["Event", "Est. UTC", "Est. PDT"],
    [
        ["Phase 1 complete + traffic scale-up (LOCUST_USERS=15)", "~15:50", "~8:50 AM"],
        ["Phase 2 complete + browser-sim start (RUM active)", "~19:10", "~12:10 PM"],
        ["Phase 3 complete (RUM data captured)", "~22:30", "~3:30 PM"],
        ["Phase 4 complete + train.jsonl export + fine-tune", "~01:00 Jul 19", "~6:00 PM"],
    ],
    col_widths=[3.8, 1.5, 1.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# 9. COST ESTIMATE
# ══════════════════════════════════════════════════════════════════════════════

heading("9. Cost Estimate", level=1)
add_table(
    ["Item", "Estimate", "Notes"],
    [
        ["AWS Bedrock (Claude Sonnet 4.6)", "$25–$45 total",
         "~$0.30–$0.47 per assessment run × ~95 remaining runs. Dominant cost."],
        ["Electricity", "< $0.30 total",
         "MacBook Pro M-series ~50–65W sustained over ~14 hrs = ~0.9 kWh"],
        ["Splunk Observability Cloud", "$0 incremental",
         "Uses existing org/entitlements; no additional data volume charges expected"],
    ],
    col_widths=[2.2, 1.5, 3.3]
)


# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/mbui/Documents/autonomous-o11y-agent/deploy/pipeline_overview.docx"
doc.save(out)
print(f"Saved: {out}")
