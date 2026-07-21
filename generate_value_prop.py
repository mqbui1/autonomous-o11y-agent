"""Generate the Autonomous O11y Agent value proposition DOCX."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Color palette (Splunk brand) ──────────────────────────────────────────────
SPLUNK_BLACK   = RGBColor(0x0d, 0x10, 0x17)   # near-black bg
SPLUNK_GREEN   = RGBColor(0x65, 0xc8, 0x73)   # primary accent
SPLUNK_ORANGE  = RGBColor(0xf4, 0x83, 0x1f)   # secondary accent
SPLUNK_WHITE   = RGBColor(0xf3, 0xf4, 0xf6)
DARK_GRAY      = RGBColor(0x1e, 0x25, 0x32)
MID_GRAY       = RGBColor(0x4b, 0x55, 0x63)
LIGHT_GRAY     = RGBColor(0x9c, 0xa3, 0xaf)
TABLE_HEADER   = RGBColor(0x1a, 0x20, 0x2c)
TABLE_ROW_ALT  = RGBColor(0xf0, 0xf4, 0xf8)
ACCENT_BLUE    = RGBColor(0x38, 0x8b, 0xfd)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = color or SPLUNK_GREEN
    elif level == 2:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = color or SPLUNK_BLACK
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = color or DARK_GRAY
    return p


def add_body(doc, text, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or SPLUNK_BLACK
    return p


def add_bullet(doc, text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.25)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = SPLUNK_BLACK
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = SPLUNK_BLACK
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = SPLUNK_BLACK
    return p


def add_callout(doc, title, text, bg=None, title_color=None):
    """Add a shaded callout box via a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg or TABLE_ROW_ALT)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = title_color or SPLUNK_GREEN
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = SPLUNK_BLACK
    cell.paragraphs[0].paragraph_format.left_indent = Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_two_col_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], TABLE_HEADER)
        hdr_cells[i].paragraphs[0].clear()
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = SPLUNK_WHITE
        r.font.size = Pt(9.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # Data rows
    for row_i, row_data in enumerate(rows):
        cells = table.rows[row_i + 1].cells
        bg = TABLE_ROW_ALT if row_i % 2 == 0 else RGBColor(0xff, 0xff, 0xff)
        for col_i, val in enumerate(row_data):
            set_cell_bg(cells[col_i], bg)
            cells[col_i].paragraphs[0].clear()
            p = cells[col_i].paragraphs[0]
            # Support bold prefix with | separator
            if "|" in str(val) and col_i == 0:
                parts = val.split("|", 1)
                r1 = p.add_run(parts[0])
                r1.bold = True
                r1.font.size = Pt(9.5)
                r1.font.color.rgb = SPLUNK_BLACK
                r2 = p.add_run("|" + parts[1])
                r2.font.size = Pt(9.5)
                r2.font.color.rgb = MID_GRAY
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                r.font.color.rgb = SPLUNK_BLACK
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
r = p.add_run("AUTONOMOUS O11Y AGENT")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = SPLUNK_GREEN

p2 = doc.add_paragraph()
r2 = p2.add_run("Value Proposition & Capability Overview")
r2.font.size = Pt(16)
r2.font.color.rgb = MID_GRAY

p3 = doc.add_paragraph()
r3 = p3.add_run(f"Splunk Observability Cloud  ·  {datetime.date.today().strftime('%B %Y')}")
r3.font.size = Pt(11)
r3.font.color.rgb = LIGHT_GRAY

doc.add_paragraph()

add_callout(
    doc,
    "TL;DR —",
    "The Autonomous O11y Agent is an AI-driven observability control plane that runs nine "
    "specialist agents in parallel against your Splunk Observability Cloud environment. "
    "It audits health, instrumentation quality, cardinality cost, detector coverage, log "
    "patterns, frontend experience, database dependencies, synthetic test coverage, and "
    "performs root cause analysis — then synthesizes all findings into a single prioritized "
    "action plan. It replaces hours of manual review with a fully automated, always-current "
    "assessment that any engineer can run in minutes.",
    bg=RGBColor(0xe8, 0xf5, 0xe9),
    title_color=RGBColor(0x2e, 0x7d, 0x32),
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. The Problem", level=1)

add_body(doc,
    "Modern observability platforms like Splunk Observability Cloud are extraordinarily "
    "powerful — but operating them well requires deep expertise across a dozen domains. "
    "Most organizations end up with a platform that is partially configured, under-utilized, "
    "and silently degrading. The symptoms are familiar:"
)

for bullet in [
    ("Detectors that never fire ", "— alerting on metrics that stopped reporting months ago (ghost detectors)"),
    ("Silent services ", "— instrumented code that stopped sending telemetry with no automated notification"),
    ("Cardinality explosions ", "— a single high-cardinality tag quietly inflating metric costs by 10×"),
    ("Instrumentation gaps ", "— missing deployment.environment breaks Related Content, APM↔Logs correlation, and Service Centric view"),
    ("No synthetic coverage ", "— services that have internal APM data but are never probed from the outside"),
    ("Database blind spots ", "— slow queries that take down services because db.system/db.name attributes were never added"),
    ("Incident investigations that take hours ", "— engineers manually correlating traces, metrics, logs, and deployment events across multiple screens"),
]:
    add_bullet(doc, bullet[1], bold_prefix=bullet[0])

doc.add_paragraph()
add_body(doc,
    "The root cause is not a lack of tools — it is a lack of automation to continuously "
    "validate that those tools are correctly configured, fully utilized, and actually "
    "catching problems. A platform review that should happen weekly gets done quarterly "
    "at best, and only when someone has the bandwidth."
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. THE SOLUTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. The Solution", level=1)

add_body(doc,
    "The Autonomous O11y Agent is an AI-powered assessment and remediation engine built "
    "specifically for Splunk Observability Cloud. It operates in two complementary modes:"
)

add_heading(doc, "Batch Mode — Periodic Deep Assessment", level=2, space_before=10)
add_body(doc,
    "Nine specialist AI agents run in parallel, each with deep expertise in a specific "
    "domain. Each specialist calls live Splunk APIs — not static rules — to understand "
    "the actual state of your environment. Findings are synthesized into a single "
    "prioritized assessment with specific service names, metric values, and recommended "
    "actions. Runs on a configurable schedule (default: every 60 minutes) or on demand."
)

add_heading(doc, "Streaming Mode — Real-Time Gateway Co-Processor", level=2, space_before=10)
add_body(doc,
    "The agent deploys alongside your OTel Collector gateway and receives a live copy "
    "of every trace and metric as it flows through — without ever blocking the primary "
    "Splunk export path. Four real-time detectors run on each data point: PII/credit card "
    "scanner, attribute validator, cardinality tracker, and new service detector. Alerts "
    "fire within seconds of a violation, and the observation buffer enriches the next "
    "batch assessment with everything the live stream detected."
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. THE NINE SPECIALISTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Nine Specialist Agents", level=1)

add_body(doc,
    "All nine specialists run in parallel (ThreadPoolExecutor, 900s timeout each). "
    "Each specialist has access to domain-specific tools, calls live Splunk APIs, "
    "and returns structured findings — not freeform text. The coordinator then performs "
    "cross-domain analysis (services flagged by multiple specialists) before a final "
    "synthesis pass with full tool access."
)
doc.add_paragraph()

specialists = [
    (
        "1 · Health",
        "agents/health.py",
        "Detectors, APM coverage, OTel Collector pipeline, license utilization",
        [
            "Audits every detector: ghost detectors (firing on dead MTS), never-fired, "
            "noisy (alert fatigue), muted rules, inactive destinations",
            "Checks APM service coverage: silent services, health check span pollution, "
            "sensitive data exposure, orphan services",
            "Verifies OTel Collector pipeline health: version currency, exporter errors, "
            "drop rates, stopped collectors",
            "Reviews license utilization headroom: APM hosts, MTS, RUM sessions, Synthetics runs",
        ],
        "Ensures the foundation of your observability platform is sound — "
        "alerts actually fire, services actually report, and you're not about to hit a license wall.",
    ),
    (
        "2 · Instrumentation",
        "agents/instrumentation.py",
        "Span quality, attribute coverage, signal completeness",
        [
            "Analyzes APM spans for missing service.name, deployment.environment, "
            "host.name, and k8s resource attributes",
            "Checks metric coverage: which services emit traces, metrics, and logs",
            "Identifies services with broken Related Content, APM↔Log correlation, "
            "and Service Centric view due to missing attributes",
            "Scores overall instrumentation quality 0–100 per service",
        ],
        "Broken Related Content is the #1 user complaint in Splunk Observability. "
        "This specialist finds and fixes the missing attributes that cause it.",
    ),
    (
        "3 · Governance",
        "agents/governance.py",
        "Metric cardinality, cost optimization, trace volume",
        [
            "Scans for metric cardinality explosions with MTS counts and cost estimates",
            "Detects slow-burn anomalies: metrics growing faster than their 7-day baseline",
            "Generates ready-to-paste OTel Collector YAML drop rules per offending dimension",
            "Snapshots trace volume per service to detect unexpected spikes",
        ],
        "A single engineer adding a user_id tag to a metric can cost tens of thousands "
        "of dollars per month in MTS overage. This specialist catches it early.",
    ),
    (
        "4 · Detector",
        "agents/detector.py",
        "Alert coverage, baseline learning, auto-provisioning",
        [
            "Discovers services with zero detector coverage ('dark services')",
            "Learns behavioral baselines from live telemetry: p50/p95/p99, error rates, request rates",
            "Provisions best-practice detectors tuned to actual traffic patterns",
            "Retunes existing detectors when baselines have drifted",
            "Auto-detects GenAI/agentic services and applies specialized detector templates",
            "Logs every deployed/retuned detector in actions_taken for full audit trail",
        ],
        "Most Splunk deployments have dark services — code running in production with no "
        "alert coverage. This specialist closes that gap automatically.",
    ),
    (
        "5 · Logs",
        "agents/logs.py",
        "Error log analysis, pattern detection, volume anomalies",
        [
            "Searches for ERROR/CRITICAL log entries across all services",
            "Groups log messages by fingerprint to surface top recurring error patterns",
            "Identifies services with zero log output (complete logging gap)",
            "Flags log volume anomalies: services generating disproportionate traffic",
            "Correlates log error counts with APM error rates to distinguish real failures "
            "from instrumentation gaps",
        ],
        "Log analysis is the most time-consuming part of any incident. "
        "This specialist pre-digests the noise into the top 5 patterns that actually matter.",
    ),
    (
        "6 · RUM",
        "agents/rum.py",
        "Frontend UX, Core Web Vitals, JavaScript errors",
        [
            "Discovers all RUM-instrumented applications and reports session volumes",
            "Assesses Core Web Vitals: LCP, FID, CLS — with pass/fail against Google thresholds",
            "Surfaces the top recurring JavaScript error types and their frequency",
            "Identifies frontend services with no RUM instrumentation despite user traffic",
            "Distinguishes instrumentation gaps from real UX degradation",
        ],
        "APM tells you what happened server-side. RUM tells you what the user actually experienced. "
        "This specialist bridges that gap and surfaces customer-impacting issues.",
    ),
    (
        "7 · RCA",
        "agents/rca.py",
        "Incident root cause analysis, causal chain investigation",
        [
            "Discovers active incidents and performs end-to-end root cause analysis",
            "Searches for error traces around the incident start time via APM GraphQL async search",
            "Analyzes topLatencyContributors per trace to pinpoint the failing operation",
            "Maps service dependency blast radius: upstream callers and downstream dependencies",
            "Correlates deployment/change events in the 90-minute window before the incident",
            "Checks Kubernetes pod CPU/memory for resource exhaustion as a causal factor",
            "Produces a causal chain with confidence: HIGH / MEDIUM / LOW",
            "Can be triggered directly from streaming alerts via run_incident_rca()",
        ],
        "Reduces mean time to root cause from hours to minutes. Instead of manually "
        "correlating traces, metrics, logs, and deployment history across multiple screens, "
        "the agent does it in one automated investigation cycle.",
    ),
    (
        "8 · Synthetics",
        "agents/synthetics.py",
        "External health validation, test coverage, performance trends",
        [
            "Inventories all Splunk Synthetics tests: browser, API, and uptime checks",
            "Identifies services with no external health validation (coverage gaps)",
            "Surfaces currently failing tests with specific error messages and per-location breakdown",
            "Detects tests with a degrading performance trend — getting slower before they fail",
            "Flags misconfigured tests: inactive, too-infrequent, or single-location only",
            "Computes uptime percentage per test over configurable windows",
        ],
        "APM tells you what happened after a user hit your service. "
        "Synthetics tells you if the service is reachable at all. "
        "This specialist ensures every critical endpoint has external validation.",
    ),
    (
        "9 · Database / Dependency",
        "agents/db.py",
        "Inferred service blind spots, DB instrumentation, slow queries",
        [
            "Maps the full service topology including inferred (unmonitored) service nodes — "
            "databases and external APIs that services call but have no instrumentation",
            "Checks db.system, db.name, and db.operation attribute coverage — required for "
            "APM Database Overview and query-level visibility",
            "Proactively surfaces the slowest outbound calls (DB queries, external APIs) "
            "before they cause incidents",
            "Detects per-service outbound error rates: high client-span errors indicate "
            "a dependency is degrading, not the calling service",
        ],
        "Most latency incidents originate in dependencies — a slow query, a rate-limited "
        "external API, a saturated connection pool. This specialist makes those dependencies "
        "visible before they cause user-facing outages.",
    ),
]

for spec in specialists:
    name, file, subtitle, bullets, value = spec

    # Specialist heading
    add_heading(doc, name, level=2, space_before=14, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{file}  ·  {subtitle}")
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = MID_GRAY

    for b in bullets:
        add_bullet(doc, b)

    # Value callout
    doc.add_paragraph()
    add_callout(doc, "Business value:", value,
                bg=RGBColor(0xf0, 0xf7, 0xff),
                title_color=ACCENT_BLUE)


# ══════════════════════════════════════════════════════════════════════════════
# 4. STREAMING MODE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Streaming Mode — Real-Time Telemetry Co-Processor", level=1)

add_body(doc,
    "In streaming mode the agent deploys as an always-on Kubernetes Deployment alongside "
    "your OTel Collector gateway. The gateway fans a copy of every trace and metric to "
    "the agent's OTLP/HTTP receiver on port 4318. The agent is never on the critical path "
    "— the secondary exporter uses retry_on_failure: false and timeout: 5s."
)
doc.add_paragraph()

add_two_col_table(doc,
    ["Streaming Detector", "What It Catches", "Alert Severity"],
    [
        ["PII / PCI Scanner",
         "Credit card numbers, SSNs, email addresses, phone numbers in span attributes",
         "CRITICAL — fires immediately"],
        ["Attribute Validator",
         "Spans missing deployment.environment, host.name, k8s.pod.name",
         "HIGH — per service"],
        ["Cardinality Tracker",
         "Sliding-window unique dimension combo counter. Warn@10K combos, critical@50K",
         "HIGH / CRITICAL"],
        ["Service Tracker",
         "New service.name appearing for the first time — triggers auto detector provisioning",
         "INFO + auto-action"],
    ],
    col_widths=[1.8, 3.5, 1.8]
)

add_body(doc,
    "All streaming observations are written into a 2-hour ObservationBuffer. When the "
    "next batch assessment runs, this buffer is injected into every specialist's context "
    "so batch findings are enriched with what the live stream detected: "
    "\"PII found in payment-service at 14:32, new service fraud-scorer appeared at 14:45.\""
)

doc.add_paragraph()
add_two_col_table(doc,
    ["Mode", "Trigger", "Latency", "Best for"],
    [
        ["Batch", "Scheduled (every N min)", "Minutes", "Deep audits, cardinality scans, detector provisioning"],
        ["Streaming", "Every span/metric", "Seconds", "PII detection, new service alerts, attribute drift"],
        ["Both (default)", "Continuous + scheduled", "Seconds + minutes", "Production: real-time + deep assessment"],
    ],
    col_widths=[1.2, 2.0, 1.1, 2.8]
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. CROSS-DOMAIN SYNTHESIS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Cross-Domain Synthesis", level=1)

add_body(doc,
    "After all nine specialists complete, the coordinator performs two additional passes "
    "before producing the final report:"
)

add_heading(doc, "Cross-Domain Analysis", level=3, space_before=10)
add_body(doc,
    "Services and issues that appear in multiple specialist domains are automatically "
    "surfaced as the highest-priority findings. A service flagged by health (silent), "
    "instrumentation (missing attributes), AND detector (no coverage) is far more "
    "critical than one flagged by a single specialist — and this is explicitly called "
    "out at the top of every report."
)

add_heading(doc, "Synthesis LLM Pass", level=3, space_before=10)
add_body(doc,
    "A final LLM pass with access to all tools across all nine specialists produces "
    "the executive summary. It can call additional tools to drill into specific "
    "cross-cutting issues that specialists surfaced but did not fully resolve — "
    "for example, calling get_trace_analysis on a service that the health, logs, "
    "AND RCA specialists all flagged."
)

add_heading(doc, "Persistent Memory & Trend Context", level=3, space_before=10)
add_body(doc,
    "Every run persists a RunRecord with active/silent service names, deployed detector IDs, "
    "critical issues, and actions taken. On the next run, trend context is injected into "
    "every specialist's prompt: services that have been silent for 2+ consecutive runs "
    "are flagged as likely instrumentation failures, not just transient gaps."
)


# ══════════════════════════════════════════════════════════════════════════════
# 6. DEPLOYMENT & INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Deployment & Integration", level=1)

add_heading(doc, "Kubernetes (Helm)", level=2, space_before=10)
add_body(doc, "A production-ready Helm chart deploys the agent in either streaming or batch mode:")
for b in [
    "Streaming mode → always-on Deployment with OTLP/HTTP receiver on port 4318",
    "Batch mode → CronJob on a configurable schedule",
    "Automatic gateway patch: one helm upgrade command fans out existing Splunk OTel Collector to the agent",
    "Persistent state volume for cross-run trend context",
    "Kubernetes Secret for Splunk + AWS credentials",
]:
    add_bullet(doc, b)

add_heading(doc, "LLM Provider", level=2, space_before=10)
add_body(doc, "The agent abstracts the LLM behind a provider interface — no lock-in:")
add_two_col_table(doc,
    ["Provider", "How to configure"],
    [
        ["AWS Bedrock (default)", "Uses boto3 Converse API. Requires AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY"],
        ["Anthropic Claude (direct)", "Set LLM_PROVIDER=anthropic, ANTHROPIC_API_KEY"],
        ["Galileo Luna", "Set LLM_PROVIDER=openai, OPENAI_BASE_URL=http://luna-host/v1"],
        ["Azure OpenAI", "Set LLM_PROVIDER=openai, OPENAI_BASE_URL=https://resource.openai.azure.com/..."],
        ["Ollama (local)", "Set LLM_PROVIDER=openai, OPENAI_BASE_URL=http://localhost:11434/v1"],
    ],
    col_widths=[2.2, 4.8]
)

add_heading(doc, "Approval Workflow", level=2, space_before=10)
add_body(doc,
    "In dry-run mode, HIGH and CRITICAL issues become numbered PendingAction items. "
    "Three approval modes are supported:"
)
for b in [
    ("Interactive: ", "stdin prompt — engineers review and approve each action"),
    ("Webhook: ", "POST {\"approved\": [1,3]} to APPROVAL_WEBHOOK_URL — integrates with Slack, PagerDuty, ServiceNow"),
    ("Auto: ", "non-interactive environments apply safe fixes automatically"),
]:
    add_bullet(doc, b[1], bold_prefix=b[0])

add_heading(doc, "Agent Self-Observability", level=2, space_before=10)
add_body(doc,
    "The agent instruments its own operations with OTel SDK, emitting spans and metrics "
    "to Splunk Observability Cloud. Build dashboards and detectors on the agent itself:"
)
for b in [
    "o11y_agent.run.duration — histogram of assessment duration",
    "o11y_agent.issues.found — counter by severity (critical/high/medium/low)",
    "o11y_agent.instrumentation_score — gauge tracking quality trend over time",
    "o11y_agent.silent_services — gauge of services with no telemetry",
]:
    add_bullet(doc, b)

add_heading(doc, "Splunk OTel Supervisor UI Integration", level=2, space_before=10)
add_body(doc,
    "The agent is designed to integrate with the Splunk OTel Supervisor — a UI-driven "
    "observability control plane. The agent exposes a REST API that the Supervisor's job "
    "runner calls as an o11y_assessment job type. Assessment findings flow into the "
    "Supervisor's recommendation/approval panel, and the Supervisor's chat interface "
    "can call the agent's tool-use loop to answer questions with real-time Splunk API data "
    "rather than static context."
)


# ══════════════════════════════════════════════════════════════════════════════
# 7. BUSINESS VALUE & ROI
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "7. Business Value & ROI", level=1)

add_two_col_table(doc,
    ["Challenge", "What the Agent Does", "Business Impact"],
    [
        ["Ghost/noisy detectors waste on-call time",
         "Health specialist audits every detector; surfaces ghost, never-fired, and high-frequency detectors with specific names",
         "Reduce alert fatigue and on-call burnout"],
        ["Silent services go unnoticed for weeks",
         "Multi-run trend context flags services silent for 2+ consecutive runs as likely instrumentation failures",
         "Catch instrumentation regressions before they become blind spots"],
        ["Cardinality explosions cause unexpected cost spikes",
         "Governance specialist detects MTS growth anomalies and generates ready-to-paste OTel Collector drop rules",
         "Prevent unbounded cost growth from a single bad metric tag"],
        ["RCA takes hours of manual correlation",
         "RCA specialist correlates traces, topology, change events, metrics, and infra in one automated cycle",
         "Reduce MTTR from hours to minutes"],
        ["New services ship with no alert coverage",
         "Service Tracker (streaming) detects new service.name in real time and triggers automatic detector provisioning",
         "Zero gap between code shipping and monitoring being active"],
        ["PII accidentally logged in spans",
         "PII Scanner (streaming) detects credit card numbers, SSNs, email in span attributes within seconds",
         "Prevent compliance violations before data reaches Splunk's index"],
        ["DB slowness invisible until outage",
         "DB/Dependency specialist surfaces slow outbound calls and missing db.* attributes proactively",
         "Catch slow queries as a trend, not as an incident"],
        ["Platform health reviews happen quarterly at best",
         "Automated assessment runs every 60 minutes with persistent trend context across runs",
         "Continuous validation — not a point-in-time snapshot"],
    ],
    col_widths=[1.9, 2.7, 2.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# 8. QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. Quick Reference", level=1)

add_heading(doc, "Key Commands", level=2, space_before=10)

add_two_col_table(doc,
    ["Command", "What it does"],
    [
        ["python3 main.py --realm us1 --token $TOKEN --environment prod",
         "One-shot full assessment, dry-run (no changes)"],
        ["python3 main.py ... --auto-apply",
         "Full assessment with auto-apply of safe fixes"],
        ["python3 main.py ... --watch --interval 30",
         "Continuous watch mode every 30 minutes"],
        ["python3 main.py ... --streaming",
         "OTLP receiver + batch assessments every 60 min"],
        ["python3 main.py ... --streaming-only",
         "Real-time streaming alerts only, no batch"],
        ["python3 main.py ... --prompt \"Why is checkout latency spiking?\"",
         "Ask the agent a specific question"],
        ["helm install o11y-agent ./charts/o11y-agent ...",
         "Deploy to Kubernetes (streaming mode by default)"],
    ],
    col_widths=[3.4, 3.7]
)

add_heading(doc, "Required Environment Variables", level=2, space_before=10)
add_two_col_table(doc,
    ["Variable", "Description"],
    [
        ["SPLUNK_REALM", "Splunk Observability realm (e.g. us1, us0, eu0)"],
        ["SPLUNK_ACCESS_TOKEN", "Splunk API access token (ingest + API scopes)"],
        ["SPLUNK_ENVIRONMENT", "Target deployment.environment name"],
        ["AWS_ACCESS_KEY_ID / SECRET", "AWS credentials for Bedrock (or set LLM_PROVIDER=openai)"],
    ],
    col_widths=[2.5, 4.6]
)

add_heading(doc, "Sibling Projects Required", level=2, space_before=10)
add_body(doc,
    "The agent wraps four existing Splunk tools as subprocess-based tools. "
    "Clone these as siblings to the agent repo:"
)
add_two_col_table(doc,
    ["Project", "Used by"],
    [
        ["auto-detector-provisioner", "Detector specialist — baseline learning + provisioning"],
        ["o11y-usage-governance", "Governance specialist — cardinality scan + drop rules"],
        ["o11y-instrumentation-analyzer", "Instrumentation specialist — attribute coverage scoring"],
        ["splunk-o11y-health-check", "Health specialist — detector audit, APM health, collector check"],
    ],
    col_widths=[2.8, 4.3]
)


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run(
    "github.com/mqbui1/autonomous-o11y-agent  ·  "
    f"Generated {datetime.date.today().strftime('%Y-%m-%d')}"
)
r.font.size = Pt(9)
r.font.color.rgb = LIGHT_GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────────────────────
out = "Autonomous_O11y_Agent_Value_Prop.docx"
doc.save(out)
print(f"Saved: {out}")
